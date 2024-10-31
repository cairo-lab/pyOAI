import math
import numpy as np
import pandas as pd
from pathlib import Path
import random
from importlib.metadata import version
from packaging.version import Version

visit_prefixes = ['P02', 'P01', 'V00', 'V01', 'V02', 'V03', 'V04', 'V05', 'V06', 'V07', 'V08', 'V09', 'V10', 'V11']
visit_prefixes_year = ['V00', 'V01', 'V03', 'V05', 'V06', 'V07', 'V08', 'V09', 'V10', 'V11']  # Drop pre-screening and half year visits
visit_cats_year = ['1: 12-month', '3: 24-month', '5: 36-month', '6: 48-month', '7: 60-month', '8: 72-month', '9: 84-month', '10: 96-month', '11: 108-month']

visit_prefix_to_month = {'V00':'0', 'V01':'12', 'V02':'18', 'V03':'24', 'V04':'30', 'V05':'36', 'V06':'48', 'V07':'60', 'V08':'72', 'V09':'84', 'V10':'96', 'V11':'108'}
visit_prefix_to_year = {'P01':-1, 'V00':0, 'V01':1, 'V02': 1.5, 'V03':2, 'V04': 2.5, 'V05':3, 'V06':4, 'V07':5, 'V08':6, 'V09':7, 'V10':8, 'V11':9}
visit_cat_to_prefix = {'1: 12-month' : 'V01', '3: 24-month': 'V03', '5: 36-month': 'V05', '6: 48-month': 'V06', '7: 60-month': 'V07', '8: 72-month': 'V08', '9: 84-month': 'V09', '10: 96-month': 'V10', '11: 108-month': 'V11'}
visit_cat_to_year = {'0: Baseline': 0, '1: 12-month' : 1, '2: 18-month' : 1.5, '3: 24-month': 2, '4: 30-month' : 2.5, '5: 36-month': 3, '6: 48-month': 4, '7: 60-month': 5, '8: 72-month': 6, '9: 84-month': 7, '10: 96-month': 8, '11: 108-month': 9, '12: 120-month': 10}

sides_cat = ['1: Right', '2: Left']

prior_post_visits = [str(i) + ' prior' for i in range(9,0,-1)] + [str(i) + ' post' for i in range(1,10)]

#  General Utilities
        
# Given a target index level and a dictionary of what values you wish to map to new values, apply to supplied dataframe
def rename_index_values(df, dct, level=0):
    df.index = df.index.set_levels([[dct.get(item, item) for item in names] if i==level else names for i, names in enumerate(df.index.levels)])
    return df

def flip_dict(dic):
    return {v:k for k,v in dic.items()}

# Fixes the annoyances of the default Pandas value_counts() method
def value_counts(ser, hidena=False):
    tmp = ser.value_counts(dropna=hidena) # Default to showing NA count
    return tmp[tmp > 0] # Drop the 0 counts that Categorical columns trigger

# This is handy when you are calling value_counts on a series of columns and want the results to print horizontally and be in the same order
def value_counts_list(ser):
    return sorted([(str(idx), val) for idx, val  in value_counts(ser).iteritems()])


#   OAI Specific 

# Given a dataframe, column/variable name and a desired value, return the set of all patient IDs that match the desired value
def get_ids(df, variable_name, match_value):
    return set(df[df[variable_name] == match_value].index.get_level_values('ID'))

# Return a list of visits where data was collected for the given column name
def get_visits(df, col):
    tmp = df[col]
    tmp = tmp.reset_index(level='Visit')
    return list(tmp[~tmp[col].isna()]['Visit'].unique())


# Pull the side name (left vs right) out of the categorical label
# Several categoricals state what side a measurement came from in the form '1: Right'
# This returns 'Right' in the former case
def get_side_name(side):
    return side.split(':')[1].strip()


# Take a dictionary of items, where each item has a list of which columns are true for each item, and turn it into a dataframe
def sets_into_dataframe(row_set_dict):
    # row_set_dict = {row_name : [cols to be True]}
    cols = list(set([item for sublist in row_set_dict.values() for item in sublist])) # Flatten lists, reduce to set, return as list of unique items
    cols.sort()
    # Create list of True/False values for each row's inclusion in a set
    for descript, row_set in row_set_dict.items():
        row_set_dict[descript] = [True if col in row_set else False for col in cols]
    df = pd.DataFrame(row_set_dict, index=cols).T
    df = df.replace(False, '-')  # Replacing False with - improves readability
    return df


# Allows the data to be split into an arbitrary number of sets
# By default, the emitted dictionary of datasets includes 'all'
# which is the original data list 
def split_data(data_list, splits_dict, include_all=True):
    # sanity check
    total = 0
    for pct in splits_dict.values():
        total += pct
    assert math.isclose(total, 1.0), "Data splits don't add up to 100%"
        
    data_cnt = len(data_list)
    random.shuffle(data_list)

    # Fill-in splits
    set_lists = {}
    if include_all:
        set_lists = {'all': data_list}
    start, end = 0, 0
    for set_name, pct in splits_dict.items():
        start, end = end, end + int(pct * (data_cnt + 1))
        set_lists[set_name] = data_list[start:end]
    
    # Due to integer division, there may be a remaining datum, detect and handle
    if end < data_cnt:
        # Find the largest set
        max_set = ('', 0)
        for set_name, size in splits_dict.items():
            if size > max_set[1]:
                max_set = (set_name, size)
        # Add remainder item to largest set
        set_lists[max_set[0]].extend(data_list[end:])
    return set_lists


# Return a list of all values that start with a SAS '.' or are an NA in a given column
def get_column_missing_values(col):
    missing_vals = set()
    tmp = value_counts(col)
    for val in list(tmp.index):
        if isinstance(val, str) and (val[0] == '.'):
            missing_vals.add(val)
        elif pd.isna(val):
            missing_vals.add(val)
    return missing_vals


# Return a list of all values that start with a '.' or are an NA across a dataframe
def get_all_missing_values(df):
    missing_vals = set()
    if isinstance(df, pd.DataFrame):
        for col in df.columns:
            missing_vals |= get_column_missing_values(df[col])
    else: # Series
        missing_vals = get_column_missing_values(df)
    return list(missing_vals)


# Return dataframe of the value counts for each column
# To avoid the visual clutter of NaNs, display the results with: display(HTML(missing_vals_df.fillna('-').to_html()))
def get_dataframe_value_counts(df, cols):
    val_cnts = []
    for col in cols:
        val_cnts.append(value_counts(df[col]))
    return pd.DataFrame(val_cnts, index=cols).T


# Return a dataframe of the value counts for only missing variables and NAs for each column
# To avoid the visual clutter of NaNs, display the results with: display(HTML(missing_vals_df.fillna('-').to_html()))
def show_missing_val_cnts(df, cols=None):
    if cols is None:
        cols = df.columns
    cnt_df = get_dataframe_value_counts(df, cols)
    idxs = [idx for idx in cnt_df.index if (isinstance(idx, str) and not idx[0].isdigit()) or pd.isna(idx)]
    return cnt_df.loc[idxs]


# Return a dataframe where no rows have missing data
def drop_rows_with_missing_data(df):
    level_count = df.index.nlevels
    df = df.dropna() # First pass, dump NA rows to speed up remaining search
    
    # Get the indexes of all rows with one or more missing value
    incompletes = pd.MultiIndex(levels=[[]]*level_count, codes=[[]]*level_count, names=df.index.names)
    for col in df.columns:
        missing_vals = [val for val in get_all_missing_values(df[col]) if isinstance(val, str)] # Drop NAs, and NaNs from the search
        incompletes = incompletes.append(df[df[col].isin(missing_vals)].index) # 

    # How many complete entries?
    complete_idx = df.index.difference(incompletes.unique())
    return df.loc[complete_idx]

# Visualization
##########################################################################

# Monkey patch to add commas to the thousands place for integers
# See: https://stackoverflow.com/questions/24922609/formatting-thousand-separator-for-integers-in-a-pandas-dataframe

if Version(version('pandas')) < Version('2.2'):
    gaf = pd.io.formats.format.GenericArrayFormatter
else:
    gaf = pd.io.formats.format._GenericArrayFormatter

class _IntArrayFormatter(gaf):

    def _format_strings(self):
        formatter = self.formatter or (lambda x: ' {:,}'.format(x))
        fmt_values = [formatter(x) for x in self.values]
        return fmt_values

pd.io.formats.format.IntArrayFormatter = _IntArrayFormatter


# Dump the dataframe rendering to something beyond the notebook.
#  If filename is 'wiki' then it just dumps mediawiki code to the page
#  to cut and past into the wiki.
#  Otherwise the target is determined by the filename extension. 
#  e.g. .png or .docx
def render_dataframe(df, filename, font_size=8):
    if filename == 'wiki':
        dataframe_to_mediawiki(df)
    else:
        extension = Path(filename).suffix
        assert extension in ['.docx', '.svg', '.png'], 'Unsupported extension. Check rendering libraries to see if it can be added'
        if extension in ['.svg', '.png']:
            dataframe_to_image(df, filename)
        elif extension == '.docx':
            dataframe_to_word(df, filename, font_size=font_size)


# Other image file types are supported but may require more code
#   See https://github.com/dexplo/dataframe_image
def dataframe_to_image(df, filename):
    import dataframe_image as dfi
    extension = Path(filename).suffix
    if extension == '.svg':
        dfi.export(df, filename, table_conversion='matplotlib')
    else:
        dfi.export(df, filename)


# Write a pandas dataframe into a MS Word docx file as a table, complete with Jupyter styling        
def dataframe_to_word(data, filename, decimal_places = 3, font='Arial', font_size = 8):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt, Mm
    from docx.oxml.shared import OxmlElement, qn

    def shade_cells(cells, shade):
        for cell in cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcVAlign = OxmlElement("w:shd")
            tcVAlign.set(qn("w:fill"), shade)
            tcPr.append(tcVAlign)
        
    def get_row_span(sequence):
        for idx in range(len(sequence)):
            if idx +1 < len(sequence):
                yield idx, sequence[idx], sequence[idx+1]

    font_size = Pt(font_size)
    white = "#ffffff"
    gray = '#f2f2f2'
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    data = pd.DataFrame(data) # My input data is in the 2D list form
    idx_lvls = data.index.nlevels
    col_lvls = data.columns.nlevels
    table = document.add_table(rows=(data.shape[0] + col_lvls), cols=data.shape[1] + idx_lvls) # First row are table headers!
    table.allow_autofit = True
    table.autofit = True
    table.style = 'Light Shading' # This style seems to keep the spacing tighter, but then I need to fix the shading below
    table.style.font.name = font
    table.style.font.size = font_size
                 
    # Column headers
    for level in range(col_lvls, ):
        last_cell = None        
        for column, label in enumerate(data.columns.get_level_values(level)):
            if last_cell is None or label != last_cell.text:            
                cell = table.cell(level, column + idx_lvls)
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER                    
                run = paragraph.add_run(label)
                run.bold = True
            last_cell = cell
    # Fix the shading of the column headers
    for idx in range(col_lvls): # Make column headers all white
        row = table.rows[idx]
        shade_cells(row.cells, white)
            
    # Row headers
    transitions = {lvl: [] for lvl in range(idx_lvls)}
    for level in range(idx_lvls):
        last_label = None
        for row, label in enumerate(data.index.get_level_values(level)):
            if last_label is None or label != last_label:
                transitions[level].append(row + col_lvls)
                cell = table.cell(row + col_lvls, level)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = paragraph.add_run(label)
                run.bold = True
            last_label = label
        transitions[level].append(data.shape[0] + col_lvls)
    # Fix the shading of the index labels            
    for level, trans_rows in transitions.items():
        column = table.columns[level]
        for idx, start, stop in get_row_span(trans_rows):
            color = white
            if idx % 2:
                color = gray
            shade_cells(column.cells[start:stop], color)        
    
    # Table contents
    for i, column in enumerate(data):
        alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_string = '{}'
        if pd.api.types.is_integer_dtype(data[column]):
            alignment = WD_ALIGN_PARAGRAPH.RIGHT
            format_string = '{:,}'
        elif pd.api.types.is_float_dtype(data[column]):
            format_string = '{:,.' + str(decimal_places) + 'f}'            
        for row in range(data.shape[0]) :
            cell = table.cell(row + col_lvls, i + idx_lvls)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.alignment = alignment            
            if pd.api.types.is_integer_dtype(data.iat[row, i]) or pd.api.types.is_float_dtype(data[column]):
                run = paragraph.add_run(format_string.format(data.iat[row, i]))
            else:
                run = paragraph.add_run(str(data.iat[row, i]))
    # Fix the shading of the rows            
    for idx in range(data.shape[0]):
        row = table.rows[idx + col_lvls]
        color = white
        if idx % 2:
            color = gray
        shade_cells(row.cells[idx_lvls:], color)

    # Hack, just make everything narrow to prevent overspacing
    # longest text len  * font_size * fudge factor
    for col in table.columns:
        col.width = int(max([len(cell.paragraphs[0].text) for cell in col.cells]) * font_size * 0.90)
        for cell in col.cells:
            cell.width = col.width
    
    document.save(filename)
    return table


def dataframe_to_mediawiki(df):
    print('TODO')
    

# Given one or more dataframes, print the wiki markdown for each table
def print_md_table(tables):
    if isinstance(tables, list):
        for table in tables:
            print(table.to_markdown(tablefmt='mediawiki'))
    else:
        print(tables.to_markdown(tablefmt='mediawiki'))    


# Long term storage
#
# Pickle is NOT viable long term storage. Even minor changes in the Pandas version
# or Python version can give you fits trying to read the data back in. This is so
# those packages don't have to worry about backwards compatibility with every fix
# or upgrade.
#
# For Pandas dataframes or series Parquet is used. For Python native data
# Msgpack is used.
#
##########################################################################

# Save Pandas as Parquet.
# Parquet is designed to be long-term storage. That said, any storage format
# will have some compatibility issue. For parquet, they are listed here:
# https://pandas.pydata.org/docs/dev/user_guide/io.html#io-parquet
# This method warns about these issues and converts types avoid information
# loss by default.
# By using pyarrow as the intermediate engine, you get greater compatibility
# than from fastparquet.
#
# forced_conversions: should be a dict of {col_name: 'type', ...} for dataframes or string 'type' for series
def write_parquet(data, filename, compression=None, convert_idx_names=True, auto_convert_cat_cols=True, forced_conversions=None, verbose=True, test=False):
    data = data.copy(deep=True)
    assert isinstance(data, pd.DataFrame) or isinstance(data, pd.Series), 'Not a Pandas data type.'
    assert (forced_conversions == None) or isinstance(forced_conversions, dict) or isinstance(forced_conversions, str), 'Not a Pandas data type.'
    if compression:
        assert Path(filename).suffix == '.' + compression and Path(filename.stem).suffix == '.parquet' , 'Extension should be .parquet.[compression_ext]'
    else:
        assert Path(filename).suffix == '.parquet', 'Extension should be .parquet or specify a compression method.'
    
    # Index level names, if specified, must be strings.
    if data.index.nlevels == 1 and data.index.name is not None and not isinstance(data.index.name, str):  # Handle single index
        if convert_idx_names:
            data.index.name = str(data.index.name)
            if verbose:
                print('Index name {}, converted to a string'.format(data.index.name))
        elif verbose:
            print('Index name {} is not a string and will be dropped in the Parquet file.'.format(data.index.name))   
    elif data.index.nlevels > 1 and not all(isinstance(name, str) for name in data.index.names):  # Multi-index
        non_string_names = [name for name in data.index.names if not isinstance(name, str)]
        if convert_idx_names:
            if verbose:
                print('Indices named {}, converted to a string'.format(non_string_names))
            data.index.names = [str(name) for name in data.index.names]                
        elif verbose:
            print('Indices named {} are not strings and will be dropped in the Parquet file.'.format(non_string_names))

    # Implement user specified conversion
    def convert_col(col, col_type):
        if col_type == 'category':
#           return col.cat.set_categories(col.cat.categories.astype('string')) # good but fails when duplicate strings are present due to different encodings
            return col.astype(str).replace('nan',np.NaN).astype('category')
        else:
            return col.astype(col_type)
    
    if forced_conversions:
        if isinstance(forced_conversions, dict):
            for col, col_type in forced_conversions.items():
                data[col] = convert_col(data[col], col_type)
        elif isinstance(forced_conversions, str):
            data = convert_col(data, forced_conversions)
            
    # In the pyarrow engine, categorical dtypes for non-string types will serialized to parquet,
    # but as coplumns of their primitive dtype (dropping categorial status). This is fine as 
    # parquet is very efficient in storing repeated values.
    def handle_cat_col(col, verbose):
        col_cat_type = col.cat.categories.dtype
        individual_cat_types = set([type(x) for x in col.cat.categories])
        
        # Catch categoricals that are all NaN - must be first check
        if col.isna().sum() == len(col):
            if verbose:
                print('Column {} is marked as categorical but is all NaN. Will be stored as float.'.format(col.name))
            return col            

        # If everything is a string, no processing necessary
        if col_cat_type == 'string':
            return col
        
        # Figure out what we are dealing with
        if col_cat_type == 'object':
            individual_cat_types = set([type(x) for x in col.cat.categories])
            
            # How many types present?
            if len(individual_cat_types) == 1:
                col_cat_type = individual_cat_types.pop()
                if verbose and isinstance(col_cat_type, str):
                    print('Column {} is marked as categorical but will be stored as {}.'.format(col.name, col_cat_type))
            else:
                if verbose:
                    print('Categorical column {} has a mix of types ({}).'.format(col.name, individual_cat_types))
                if auto_convert_cat_cols and any(isinstance(x, str) for x in individual_cat_types):
                    if verbose:
                        print('\tConverting all to string and saving as a categorical in the Parquet file.')
                    return convert_col(data[col], 'category')
                elif verbose:
                    print('\tWill not be saved as categorical. Type will be chosen by parquet code.')
        elif verbose:
            print('Column {} is marked as categorical but will be stored as {}.'.format(col.name, col_cat_type)) 
    
        return col
        
    if isinstance(data, pd.Series) and isinstance(data.dtype, pd.CategoricalDtype):
        data = handle_cat_col(data, verbose)
    elif isinstance(data, pd.DataFrame):
        for col in data.select_dtypes(include='category'):
            data[col] = handle_cat_col(data[col], verbose)
    
    # Write to disk
    if not test:
        if isinstance(filename, Path):  # when filename is a path like 'data/my_df.parquet'
            filename.parent.mkdir(parents=True, exist_ok=True)
        # See: https://pandas.pydata.org/docs/dev/reference/api/pandas.DataFrame.to_parquet.html
        data.to_parquet(filename, engine="pyarrow", compression=compression)

        
# Read parquet file and unless specified, optimize the 
# arquet doesn't do categorical columns out of anything but strings,         
def read_parquet(filename, compress_cols=True, threshold=0.10):
    assert any([ext == '.parquet' for ext in Path(filename).suffixes]), 'File lacks .parquet extension'
    # https://pandas.pydata.org/docs/reference/api/pandas.read_parquet.html
    df = pd.read_parquet(filename, engine='pyarrow')
    
    # Convert columns with mostly the same values to categoricals. Saves memory.
    if compress_cols:
        row_cnt = len(df)
        non_cat_cols = df.select_dtypes(exclude='category')
        for col in non_cat_cols:
            idx = df[col].first_valid_index()
            
            # So...parquet can handle lists as a row, col value, but reads them in as numpy arrays
            # This makes at least one verion of Pandas choke if you run value_counts on that col
            # This catches those cases and converts them back to lists
            if idx and isinstance(df[col][idx], np.ndarray):
                df[col] = df[col].apply(lambda x: list(x) if isinstance(x, np.ndarray) else x)
                continue

            possible_cats = df[col].value_counts().index
            # What percent of columns are a unique value?
            if (len(possible_cats) / row_cnt) <= threshold:
                if any([isinstance(cat, str) for cat in possible_cats]):
                    df[col] = df[col].astype(str).replace('nan', np.NaN).astype('category')
                else:
                    df[col] = df[col].astype('category')
    df.sort_index()
    return df


# Write data to msgpack file
def write_msgpack(data, filename):
    assert Path(filename).suffix == '.msgpack', 'Expecting a file extension of .msgpack'
    assert not (isinstance(data, pd.DataFrame) or isinstance(data, pd.Series)), 'Use write_pandas() for dataframes and series.'
    if isinstance(filename, Path):
        filename.parent.mkdir(parents=True, exist_ok=True)
    with open(filename, "wb") as outfile:
        import msgpack
        import msgpack_numpy as m  # encoding/decoding routines that enable the serialization/deserialization of data types provided by numpy
        m.patch()
        packed = msgpack.packb(data)
        outfile.write(packed)


# Read either parquet or msgpack
def read_msgpack(filename):
    assert Path(filename).suffix == '.msgpack', 'Expecting a file extension of .msgpack'
    with open(filename, "rb") as data_file:
        import msgpack
        import msgpack_numpy as m  # encoding/decoding routines that enable the serialization/deserialization of data types provided by numpy
        m.patch()        
        return msgpack.unpackb(data_file.read(), strict_map_key=False)
